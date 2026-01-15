"""
Document Parser Service

Extracts text content from various document formats:
- PDF files using PyPDF2
- DOCX files using python-docx
- TXT files (plain text)
"""

import io
import logging
from typing import Optional
from pathlib import Path

from PyPDF2 import PdfReader
from PyPDF2.errors import PdfReadError
from docx import Document as DocxDocument
from docx.opc.exceptions import PackageNotFoundError
from pydantic import BaseModel

logger = logging.getLogger(__name__)


class ParsedDocument(BaseModel):
    """Result of document parsing."""
    text: str
    page_count: Optional[int] = None
    word_count: int
    char_count: int
    file_type: str
    success: bool
    error_message: Optional[str] = None


class DocumentParser:
    """
    Extract text from PDF, DOCX, and TXT documents.
    """

    SUPPORTED_TYPES = {".pdf", ".docx", ".doc", ".txt"}

    def __init__(self, max_file_size_mb: int = 50):
        self.max_file_size = max_file_size_mb * 1024 * 1024  # Convert to bytes

    def parse(self, file_content: bytes, filename: str) -> ParsedDocument:
        """
        Parse document and extract text.

        Args:
            file_content: Raw bytes of the file
            filename: Original filename (used to determine type)

        Returns:
            ParsedDocument with extracted text and metadata
        """
        # Check file size
        if len(file_content) > self.max_file_size:
            return ParsedDocument(
                text="",
                word_count=0,
                char_count=0,
                file_type="unknown",
                success=False,
                error_message=f"File exceeds maximum size of {self.max_file_size // (1024*1024)}MB",
            )

        # Determine file type
        ext = Path(filename).suffix.lower()
        if ext not in self.SUPPORTED_TYPES:
            return ParsedDocument(
                text="",
                word_count=0,
                char_count=0,
                file_type=ext,
                success=False,
                error_message=f"Unsupported file type: {ext}",
            )

        # Route to appropriate parser
        try:
            if ext == ".pdf":
                return self._parse_pdf(file_content)
            elif ext in {".docx", ".doc"}:
                return self._parse_docx(file_content)
            elif ext == ".txt":
                return self._parse_txt(file_content)
            else:
                return ParsedDocument(
                    text="",
                    word_count=0,
                    char_count=0,
                    file_type=ext,
                    success=False,
                    error_message=f"No parser available for {ext}",
                )
        except Exception as e:
            logger.error(f"Error parsing {filename}: {e}")
            return ParsedDocument(
                text="",
                word_count=0,
                char_count=0,
                file_type=ext,
                success=False,
                error_message=str(e),
            )

    def _parse_pdf(self, file_content: bytes) -> ParsedDocument:
        """Extract text from PDF file."""
        try:
            pdf_file = io.BytesIO(file_content)
            reader = PdfReader(pdf_file)

            text_parts = []
            for page_num, page in enumerate(reader.pages):
                try:
                    page_text = page.extract_text()
                    if page_text:
                        text_parts.append(f"--- Page {page_num + 1} ---\n{page_text}")
                except Exception as e:
                    logger.warning(f"Error extracting text from page {page_num + 1}: {e}")
                    text_parts.append(f"--- Page {page_num + 1} ---\n[Error extracting text]")

            full_text = "\n\n".join(text_parts)

            return ParsedDocument(
                text=full_text,
                page_count=len(reader.pages),
                word_count=len(full_text.split()),
                char_count=len(full_text),
                file_type="pdf",
                success=True,
            )

        except PdfReadError as e:
            logger.error(f"PDF read error: {e}")
            return ParsedDocument(
                text="",
                page_count=0,
                word_count=0,
                char_count=0,
                file_type="pdf",
                success=False,
                error_message=f"Failed to read PDF: {str(e)}",
            )

    def _parse_docx(self, file_content: bytes) -> ParsedDocument:
        """Extract text from DOCX file."""
        try:
            docx_file = io.BytesIO(file_content)
            doc = DocxDocument(docx_file)

            # Extract text from paragraphs
            paragraphs = []
            for para in doc.paragraphs:
                if para.text.strip():
                    paragraphs.append(para.text)

            # Also extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        paragraphs.append(" | ".join(row_text))

            full_text = "\n\n".join(paragraphs)

            return ParsedDocument(
                text=full_text,
                page_count=None,  # DOCX doesn't have fixed page counts
                word_count=len(full_text.split()),
                char_count=len(full_text),
                file_type="docx",
                success=True,
            )

        except PackageNotFoundError as e:
            logger.error(f"DOCX package error: {e}")
            return ParsedDocument(
                text="",
                word_count=0,
                char_count=0,
                file_type="docx",
                success=False,
                error_message=f"Failed to read DOCX: {str(e)}",
            )

    def _parse_txt(self, file_content: bytes) -> ParsedDocument:
        """Extract text from plain text file."""
        try:
            # Try common encodings
            encodings = ["utf-8", "utf-16", "latin-1", "cp1252"]
            text = None

            for encoding in encodings:
                try:
                    text = file_content.decode(encoding)
                    break
                except UnicodeDecodeError:
                    continue

            if text is None:
                return ParsedDocument(
                    text="",
                    word_count=0,
                    char_count=0,
                    file_type="txt",
                    success=False,
                    error_message="Could not decode text file with any supported encoding",
                )

            return ParsedDocument(
                text=text,
                page_count=None,
                word_count=len(text.split()),
                char_count=len(text),
                file_type="txt",
                success=True,
            )

        except Exception as e:
            logger.error(f"Text file parsing error: {e}")
            return ParsedDocument(
                text="",
                word_count=0,
                char_count=0,
                file_type="txt",
                success=False,
                error_message=str(e),
            )

    def parse_multiple(self, documents: list[tuple[bytes, str]]) -> list[ParsedDocument]:
        """
        Parse multiple documents and combine results.

        Args:
            documents: List of (file_content, filename) tuples

        Returns:
            List of ParsedDocument results
        """
        return [self.parse(content, filename) for content, filename in documents]

    def combine_documents(self, parsed_docs: list[ParsedDocument]) -> str:
        """
        Combine multiple parsed documents into a single text string.

        Args:
            parsed_docs: List of ParsedDocument objects

        Returns:
            Combined text with document separators
        """
        combined_parts = []
        for i, doc in enumerate(parsed_docs):
            if doc.success and doc.text:
                combined_parts.append(
                    f"=== Document {i + 1} ({doc.file_type.upper()}) ===\n{doc.text}"
                )

        return "\n\n" + "\n\n".join(combined_parts) if combined_parts else ""


# Singleton instance
_parser: Optional[DocumentParser] = None


def get_document_parser() -> DocumentParser:
    """Get or create the singleton DocumentParser instance."""
    global _parser
    if _parser is None:
        _parser = DocumentParser()
    return _parser
