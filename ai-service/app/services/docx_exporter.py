"""
Word Document Export Service

Generates professional Word documents from demand letter content.
"""

import io
import re
from typing import Optional
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Twips
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


class ExportOptions:
    """Configuration options for document export."""

    def __init__(
        self,
        font_name: str = "Times New Roman",
        font_size: int = 12,
        margin_top: float = 1.0,
        margin_bottom: float = 1.0,
        margin_left: float = 1.0,
        margin_right: float = 1.0,
        line_spacing: float = 1.0,
        include_letterhead: bool = False,
        letterhead_firm_name: Optional[str] = None,
        letterhead_address: Optional[str] = None,
        letterhead_phone: Optional[str] = None,
        letterhead_email: Optional[str] = None,
        include_page_numbers: bool = True,
        include_date: bool = True,
    ):
        self.font_name = font_name
        self.font_size = font_size
        self.margin_top = margin_top
        self.margin_bottom = margin_bottom
        self.margin_left = margin_left
        self.margin_right = margin_right
        self.line_spacing = line_spacing
        self.include_letterhead = include_letterhead
        self.letterhead_firm_name = letterhead_firm_name
        self.letterhead_address = letterhead_address
        self.letterhead_phone = letterhead_phone
        self.letterhead_email = letterhead_email
        self.include_page_numbers = include_page_numbers
        self.include_date = include_date


class DocxExporter:
    """Service for exporting demand letters to Word documents."""

    def __init__(self):
        pass

    def export(
        self,
        content: str,
        title: str,
        options: Optional[ExportOptions] = None
    ) -> bytes:
        """
        Export demand letter content to a Word document.

        Args:
            content: The demand letter text content
            title: Document title (used for metadata)
            options: Export configuration options

        Returns:
            Word document as bytes
        """
        if options is None:
            options = ExportOptions()

        # Create new document
        doc = Document()

        # Set up document metadata
        doc.core_properties.title = title
        doc.core_properties.author = "Demand Letter Generator"

        # Set up document styles
        self._setup_styles(doc, options)

        # Set page margins
        self._set_margins(doc, options)

        # Add letterhead if requested
        if options.include_letterhead:
            self._add_letterhead(doc, options)

        # Add page numbers if requested
        if options.include_page_numbers:
            self._add_page_numbers(doc)

        # Parse and add content
        self._add_content(doc, content, options)

        # Save to bytes buffer
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)

        return buffer.getvalue()

    def _setup_styles(self, doc: Document, options: ExportOptions) -> None:
        """Set up document styles."""
        # Normal style
        style = doc.styles['Normal']
        font = style.font
        font.name = options.font_name
        font.size = Pt(options.font_size)

        # Set paragraph format
        paragraph_format = style.paragraph_format
        paragraph_format.space_after = Pt(0)
        paragraph_format.space_before = Pt(0)

        if options.line_spacing == 1.0:
            paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        elif options.line_spacing == 1.5:
            paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        elif options.line_spacing == 2.0:
            paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
        else:
            paragraph_format.line_spacing = options.line_spacing

        # Heading 1 style
        try:
            h1_style = doc.styles['Heading 1']
        except KeyError:
            h1_style = doc.styles.add_style('Heading 1', WD_STYLE_TYPE.PARAGRAPH)

        h1_font = h1_style.font
        h1_font.name = options.font_name
        h1_font.size = Pt(options.font_size + 4)
        h1_font.bold = True
        h1_style.paragraph_format.space_before = Pt(12)
        h1_style.paragraph_format.space_after = Pt(6)

        # Heading 2 style
        try:
            h2_style = doc.styles['Heading 2']
        except KeyError:
            h2_style = doc.styles.add_style('Heading 2', WD_STYLE_TYPE.PARAGRAPH)

        h2_font = h2_style.font
        h2_font.name = options.font_name
        h2_font.size = Pt(options.font_size + 2)
        h2_font.bold = True
        h2_style.paragraph_format.space_before = Pt(10)
        h2_style.paragraph_format.space_after = Pt(4)

        # Letterhead style
        letterhead_style = doc.styles.add_style('Letterhead', WD_STYLE_TYPE.PARAGRAPH)
        letterhead_font = letterhead_style.font
        letterhead_font.name = options.font_name
        letterhead_font.size = Pt(options.font_size)
        letterhead_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        letterhead_style.paragraph_format.space_after = Pt(0)

    def _set_margins(self, doc: Document, options: ExportOptions) -> None:
        """Set page margins."""
        for section in doc.sections:
            section.top_margin = Inches(options.margin_top)
            section.bottom_margin = Inches(options.margin_bottom)
            section.left_margin = Inches(options.margin_left)
            section.right_margin = Inches(options.margin_right)

    def _add_letterhead(self, doc: Document, options: ExportOptions) -> None:
        """Add letterhead to document."""
        if options.letterhead_firm_name:
            p = doc.add_paragraph(options.letterhead_firm_name, style='Letterhead')
            run = p.runs[0]
            run.bold = True
            run.font.size = Pt(options.font_size + 2)

        if options.letterhead_address:
            p = doc.add_paragraph(options.letterhead_address, style='Letterhead')
            p.runs[0].font.size = Pt(options.font_size - 1)

        if options.letterhead_phone or options.letterhead_email:
            contact_parts = []
            if options.letterhead_phone:
                contact_parts.append(f"Tel: {options.letterhead_phone}")
            if options.letterhead_email:
                contact_parts.append(f"Email: {options.letterhead_email}")
            p = doc.add_paragraph(" | ".join(contact_parts), style='Letterhead')
            p.runs[0].font.size = Pt(options.font_size - 1)

        # Add horizontal line
        self._add_horizontal_line(doc)

        # Add some space after letterhead
        doc.add_paragraph()

    def _add_horizontal_line(self, doc: Document) -> None:
        """Add a horizontal line."""
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(6)

        # Create bottom border
        pPr = p._p.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'), 'single')
        bottom.set(qn('w:sz'), '6')
        bottom.set(qn('w:space'), '1')
        bottom.set(qn('w:color'), '000000')
        pBdr.append(bottom)
        pPr.append(pBdr)

    def _add_page_numbers(self, doc: Document) -> None:
        """Add page numbers to document footer."""
        for section in doc.sections:
            footer = section.footer
            footer.is_linked_to_previous = False

            p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER

            # Add page number field
            run = p.add_run()
            fldChar1 = OxmlElement('w:fldChar')
            fldChar1.set(qn('w:fldCharType'), 'begin')
            run._r.append(fldChar1)

            run2 = p.add_run()
            instrText = OxmlElement('w:instrText')
            instrText.text = "PAGE"
            run2._r.append(instrText)

            run3 = p.add_run()
            fldChar2 = OxmlElement('w:fldChar')
            fldChar2.set(qn('w:fldCharType'), 'end')
            run3._r.append(fldChar2)

    def _add_content(self, doc: Document, content: str, options: ExportOptions) -> None:
        """Parse and add content to document."""
        # Split content into lines
        lines = content.split('\n')

        i = 0
        while i < len(lines):
            line = lines[i]
            stripped = line.strip()

            # Skip empty lines but preserve paragraph breaks
            if not stripped:
                doc.add_paragraph()
                i += 1
                continue

            # Check for markdown-style headers
            if stripped.startswith('# '):
                p = doc.add_paragraph(stripped[2:], style='Heading 1')
                i += 1
                continue
            elif stripped.startswith('## '):
                p = doc.add_paragraph(stripped[3:], style='Heading 2')
                i += 1
                continue
            elif stripped.startswith('### '):
                p = doc.add_paragraph(stripped[4:])
                run = p.runs[0]
                run.bold = True
                i += 1
                continue

            # Check for bullet points
            if stripped.startswith('- ') or stripped.startswith('* '):
                p = doc.add_paragraph(stripped[2:], style='List Bullet')
                i += 1
                continue

            # Check for numbered lists
            numbered_match = re.match(r'^(\d+)\.\s+(.+)$', stripped)
            if numbered_match:
                p = doc.add_paragraph(numbered_match.group(2), style='List Number')
                i += 1
                continue

            # Check for bold text markers
            if stripped.startswith('**') and stripped.endswith('**'):
                p = doc.add_paragraph()
                run = p.add_run(stripped[2:-2])
                run.bold = True
                i += 1
                continue

            # Regular paragraph - handle inline formatting
            p = doc.add_paragraph()
            self._add_formatted_text(p, stripped)
            i += 1

    def _add_formatted_text(self, paragraph, text: str) -> None:
        """Add text with inline formatting (bold, italic, underline)."""
        # Pattern to match **bold**, *italic*, and __underline__
        pattern = r'(\*\*[^*]+\*\*|\*[^*]+\*|__[^_]+__)'
        parts = re.split(pattern, text)

        for part in parts:
            if not part:
                continue

            if part.startswith('**') and part.endswith('**'):
                # Bold
                run = paragraph.add_run(part[2:-2])
                run.bold = True
            elif part.startswith('*') and part.endswith('*'):
                # Italic
                run = paragraph.add_run(part[1:-1])
                run.italic = True
            elif part.startswith('__') and part.endswith('__'):
                # Underline
                run = paragraph.add_run(part[2:-2])
                run.underline = True
            else:
                # Regular text
                paragraph.add_run(part)


# Singleton instance
_exporter_instance: Optional[DocxExporter] = None


def get_docx_exporter() -> DocxExporter:
    """Get the DocxExporter singleton instance."""
    global _exporter_instance
    if _exporter_instance is None:
        _exporter_instance = DocxExporter()
    return _exporter_instance
