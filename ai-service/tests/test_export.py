"""
Tests for the Word document export functionality.
"""

import pytest
import io
import zipfile
from fastapi.testclient import TestClient
from docx import Document

from app.main import app
from app.services.docx_exporter import DocxExporter, ExportOptions


client = TestClient(app)


class TestDocxExporter:
    """Tests for the DocxExporter service."""

    def test_basic_export(self):
        """Test basic document export."""
        exporter = DocxExporter()
        content = "This is a test demand letter.\n\nSecond paragraph."
        title = "Test Letter"

        result = exporter.export(content, title)

        assert isinstance(result, bytes)
        assert len(result) > 0

        # Verify it's a valid docx
        doc = Document(io.BytesIO(result))
        assert doc.core_properties.title == title

    def test_export_with_headers(self):
        """Test export with markdown-style headers."""
        exporter = DocxExporter()
        content = """# Main Header

This is a paragraph.

## Sub Header

Another paragraph.

### Small Header

Final paragraph."""

        result = exporter.export(content, "Headers Test")

        doc = Document(io.BytesIO(result))
        paragraphs = list(doc.paragraphs)
        assert len(paragraphs) > 0

    def test_export_with_bullet_points(self):
        """Test export with bullet points."""
        exporter = DocxExporter()
        content = """Introduction paragraph.

- First bullet point
- Second bullet point
- Third bullet point

Closing paragraph."""

        result = exporter.export(content, "Bullets Test")

        doc = Document(io.BytesIO(result))
        assert len(doc.paragraphs) > 0

    def test_export_with_numbered_list(self):
        """Test export with numbered list."""
        exporter = DocxExporter()
        content = """Steps to follow:

1. First step
2. Second step
3. Third step

Done."""

        result = exporter.export(content, "Numbered List Test")

        doc = Document(io.BytesIO(result))
        assert len(doc.paragraphs) > 0

    def test_export_with_bold_text(self):
        """Test export with bold text markers."""
        exporter = DocxExporter()
        content = """This is a **bold statement** in the letter.

**Complete Bold Line**

Normal line with **multiple** bold **words**."""

        result = exporter.export(content, "Bold Test")

        doc = Document(io.BytesIO(result))
        assert len(doc.paragraphs) > 0

    def test_export_with_custom_font(self):
        """Test export with custom font options."""
        exporter = DocxExporter()
        options = ExportOptions(
            font_name="Arial",
            font_size=14
        )

        result = exporter.export("Test content", "Font Test", options)

        doc = Document(io.BytesIO(result))
        # Verify document was created with custom settings
        style = doc.styles['Normal']
        assert style.font.name == "Arial"
        assert style.font.size.pt == 14

    def test_export_with_margins(self):
        """Test export with custom margins."""
        exporter = DocxExporter()
        options = ExportOptions(
            margin_top=1.5,
            margin_bottom=1.5,
            margin_left=1.25,
            margin_right=1.25
        )

        result = exporter.export("Test content", "Margins Test", options)

        doc = Document(io.BytesIO(result))
        section = doc.sections[0]
        # Verify margins (1 inch = 914400 EMUs)
        assert abs(section.top_margin.inches - 1.5) < 0.01
        assert abs(section.left_margin.inches - 1.25) < 0.01

    def test_export_with_letterhead(self):
        """Test export with letterhead."""
        exporter = DocxExporter()
        options = ExportOptions(
            include_letterhead=True,
            letterhead_firm_name="Test Law Firm, LLP",
            letterhead_address="123 Legal Street, Suite 100, New York, NY 10001",
            letterhead_phone="(555) 123-4567",
            letterhead_email="info@testfirm.com"
        )

        result = exporter.export("Test content", "Letterhead Test", options)

        doc = Document(io.BytesIO(result))
        paragraphs = list(doc.paragraphs)
        # Should have letterhead paragraphs
        assert len(paragraphs) > 3

    def test_export_with_page_numbers(self):
        """Test export with page numbers enabled."""
        exporter = DocxExporter()
        options = ExportOptions(include_page_numbers=True)

        result = exporter.export("Test content", "Page Numbers Test", options)

        doc = Document(io.BytesIO(result))
        # Verify footer exists
        section = doc.sections[0]
        assert section.footer is not None


class TestExportAPI:
    """Tests for the export API endpoints."""

    def test_export_endpoint(self):
        """Test single document export endpoint."""
        response = client.post(
            "/ai/export",
            json={
                "content": "This is a test demand letter for export.",
                "title": "Test Export"
            }
        )

        assert response.status_code == 200
        assert response.headers["content-type"] == "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        assert "content-disposition" in response.headers
        assert len(response.content) > 0

    def test_export_with_options(self):
        """Test export with custom options."""
        response = client.post(
            "/ai/export",
            json={
                "content": "Test content with options.",
                "title": "Options Test",
                "options": {
                    "font_name": "Arial",
                    "font_size": 11,
                    "include_page_numbers": True,
                    "line_spacing": 1.5
                }
            }
        )

        assert response.status_code == 200
        assert len(response.content) > 0

    def test_export_with_letterhead(self):
        """Test export with letterhead options."""
        response = client.post(
            "/ai/export",
            json={
                "content": "Test content with letterhead.",
                "title": "Letterhead Export Test",
                "options": {
                    "include_letterhead": True,
                    "letterhead_firm_name": "API Test Firm",
                    "letterhead_address": "456 Test Ave"
                }
            }
        )

        assert response.status_code == 200
        assert len(response.content) > 0

    def test_export_empty_content(self):
        """Test export with empty content."""
        response = client.post(
            "/ai/export",
            json={
                "content": "",
                "title": "Empty Test"
            }
        )

        # Should still succeed with empty content
        assert response.status_code == 200

    def test_batch_export_endpoint(self):
        """Test batch export endpoint."""
        response = client.post(
            "/ai/export/batch",
            json={
                "items": [
                    {"id": "1", "content": "First letter content", "title": "Letter One"},
                    {"id": "2", "content": "Second letter content", "title": "Letter Two"},
                    {"id": "3", "content": "Third letter content", "title": "Letter Three"}
                ]
            }
        )

        assert response.status_code == 200
        assert response.headers["content-type"] == "application/zip"

        # Verify ZIP contents
        zip_file = zipfile.ZipFile(io.BytesIO(response.content))
        names = zip_file.namelist()
        assert len(names) == 3
        assert all(name.endswith(".docx") for name in names)

    def test_batch_export_with_options(self):
        """Test batch export with shared options."""
        response = client.post(
            "/ai/export/batch",
            json={
                "items": [
                    {"id": "1", "content": "Content 1", "title": "Title 1"},
                    {"id": "2", "content": "Content 2", "title": "Title 2"}
                ],
                "options": {
                    "font_name": "Georgia",
                    "font_size": 12
                }
            }
        )

        assert response.status_code == 200
        assert response.headers["x-export-file-count"] == "2"

    def test_batch_export_empty_items(self):
        """Test batch export with empty items list."""
        response = client.post(
            "/ai/export/batch",
            json={"items": []}
        )

        assert response.status_code == 400

    def test_batch_export_max_items(self):
        """Test batch export exceeding max items."""
        items = [
            {"id": str(i), "content": f"Content {i}", "title": f"Title {i}"}
            for i in range(51)
        ]

        response = client.post(
            "/ai/export/batch",
            json={"items": items}
        )

        assert response.status_code == 400

    def test_batch_export_with_custom_filenames(self):
        """Test batch export with custom filenames."""
        response = client.post(
            "/ai/export/batch",
            json={
                "items": [
                    {"id": "1", "content": "Content 1", "title": "Title 1", "filename": "custom_name_1"},
                    {"id": "2", "content": "Content 2", "title": "Title 2", "filename": "custom_name_2"}
                ]
            }
        )

        assert response.status_code == 200

        zip_file = zipfile.ZipFile(io.BytesIO(response.content))
        names = zip_file.namelist()
        assert "custom_name_1.docx" in names
        assert "custom_name_2.docx" in names

    def test_export_options_endpoint(self):
        """Test export options endpoint."""
        response = client.get("/ai/export/options")

        assert response.status_code == 200
        data = response.json()
        assert "fonts" in data
        assert "defaults" in data
        assert "font_sizes" in data
        assert "line_spacing_options" in data
        assert "margin_range" in data


class TestExportEdgeCases:
    """Tests for edge cases in export functionality."""

    def test_export_with_special_characters(self):
        """Test export with special characters in content."""
        exporter = DocxExporter()
        content = """Special characters test:

& ampersand
< less than
> greater than
" double quote
' single quote
© copyright
® registered
™ trademark"""

        result = exporter.export(content, "Special Chars Test")
        assert isinstance(result, bytes)
        assert len(result) > 0

    def test_export_with_unicode(self):
        """Test export with unicode characters."""
        exporter = DocxExporter()
        content = """Unicode test:

Chinese: 你好世界
Japanese: こんにちは
Korean: 안녕하세요
Arabic: مرحبا
Emoji: Test 🎉"""

        result = exporter.export(content, "Unicode Test")
        assert isinstance(result, bytes)
        assert len(result) > 0

    def test_export_long_content(self):
        """Test export with very long content."""
        exporter = DocxExporter()
        # Generate ~100KB of content
        content = "Lorem ipsum dolor sit amet. " * 5000

        result = exporter.export(content, "Long Content Test")
        assert isinstance(result, bytes)
        assert len(result) > 0

    def test_export_special_title(self):
        """Test export with special characters in title."""
        response = client.post(
            "/ai/export",
            json={
                "content": "Test content",
                "title": "Letter: Smith v. Jones (Case #123)"
            }
        )

        assert response.status_code == 200
        # Filename should be sanitized
        disposition = response.headers["content-disposition"]
        assert "Smith" in disposition

    def test_batch_export_duplicate_titles(self):
        """Test batch export with duplicate titles."""
        response = client.post(
            "/ai/export/batch",
            json={
                "items": [
                    {"id": "1", "content": "Content 1", "title": "Same Title"},
                    {"id": "2", "content": "Content 2", "title": "Same Title"},
                    {"id": "3", "content": "Content 3", "title": "Same Title"}
                ]
            }
        )

        assert response.status_code == 200

        zip_file = zipfile.ZipFile(io.BytesIO(response.content))
        names = zip_file.namelist()
        # Should have unique filenames
        assert len(names) == len(set(names))
