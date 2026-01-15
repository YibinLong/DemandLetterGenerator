"""
Unit tests for the Word document export service (independent of other services).
"""

import pytest
import io
import sys
import os

# Import docx_exporter module directly without going through __init__.py
sys.path.insert(0, os.path.join(os.path.dirname(__file__), '..', 'app', 'services'))
from docx_exporter import DocxExporter, ExportOptions, get_docx_exporter
from docx import Document


class TestDocxExporter:
    """Tests for the DocxExporter service."""

    def test_basic_export(self):
        """Test basic document export."""
        exporter = DocxExporter()
        content = "This is a test demand letter.\n\nSecond paragraph."
        title = "Test Letter"

        result = exporter.export(content, title)

        assert isinstance(result, bytes)
        assert len(result) > 0

        # Verify it's a valid docx
        doc = Document(io.BytesIO(result))
        assert doc.core_properties.title == title

    def test_export_with_headers(self):
        """Test export with markdown-style headers."""
        exporter = DocxExporter()
        content = """# Main Header

This is a paragraph.

## Sub Header

Another paragraph.

### Small Header

Final paragraph."""

        result = exporter.export(content, "Headers Test")

        doc = Document(io.BytesIO(result))
        paragraphs = list(doc.paragraphs)
        assert len(paragraphs) > 0

    def test_export_with_bullet_points(self):
        """Test export with bullet points."""
        exporter = DocxExporter()
        content = """Introduction paragraph.

- First bullet point
- Second bullet point
- Third bullet point

Closing paragraph."""

        result = exporter.export(content, "Bullets Test")

        doc = Document(io.BytesIO(result))
        assert len(doc.paragraphs) > 0

    def test_export_with_numbered_list(self):
        """Test export with numbered list."""
        exporter = DocxExporter()
        content = """Steps to follow:

1. First step
2. Second step
3. Third step

Done."""

        result = exporter.export(content, "Numbered List Test")

        doc = Document(io.BytesIO(result))
        assert len(doc.paragraphs) > 0

    def test_export_with_bold_text(self):
        """Test export with bold text markers."""
        exporter = DocxExporter()
        content = """This is a **bold statement** in the letter.

**Complete Bold Line**

Normal line with **multiple** bold **words**."""

        result = exporter.export(content, "Bold Test")

        doc = Document(io.BytesIO(result))
        assert len(doc.paragraphs) > 0

    def test_export_with_custom_font(self):
        """Test export with custom font options."""
        exporter = DocxExporter()
        options = ExportOptions(
            font_name="Arial",
            font_size=14
        )

        result = exporter.export("Test content", "Font Test", options)

        doc = Document(io.BytesIO(result))
        # Verify document was created with custom settings
        style = doc.styles['Normal']
        assert style.font.name == "Arial"
        assert style.font.size.pt == 14

    def test_export_with_margins(self):
        """Test export with custom margins."""
        exporter = DocxExporter()
        options = ExportOptions(
            margin_top=1.5,
            margin_bottom=1.5,
            margin_left=1.25,
            margin_right=1.25
        )

        result = exporter.export("Test content", "Margins Test", options)

        doc = Document(io.BytesIO(result))
        section = doc.sections[0]
        # Verify margins (1 inch = 914400 EMUs)
        assert abs(section.top_margin.inches - 1.5) < 0.01
        assert abs(section.left_margin.inches - 1.25) < 0.01

    def test_export_with_letterhead(self):
        """Test export with letterhead."""
        exporter = DocxExporter()
        options = ExportOptions(
            include_letterhead=True,
            letterhead_firm_name="Test Law Firm, LLP",
            letterhead_address="123 Legal Street, Suite 100, New York, NY 10001",
            letterhead_phone="(555) 123-4567",
            letterhead_email="info@testfirm.com"
        )

        result = exporter.export("Test content", "Letterhead Test", options)

        doc = Document(io.BytesIO(result))
        paragraphs = list(doc.paragraphs)
        # Should have letterhead paragraphs
        assert len(paragraphs) > 3

    def test_export_with_page_numbers(self):
        """Test export with page numbers enabled."""
        exporter = DocxExporter()
        options = ExportOptions(include_page_numbers=True)

        result = exporter.export("Test content", "Page Numbers Test", options)

        doc = Document(io.BytesIO(result))
        # Verify footer exists
        section = doc.sections[0]
        assert section.footer is not None

    def test_singleton_instance(self):
        """Test that get_docx_exporter returns singleton."""
        exporter1 = get_docx_exporter()
        exporter2 = get_docx_exporter()

        assert exporter1 is exporter2


class TestExportEdgeCases:
    """Tests for edge cases in export functionality."""

    def test_export_with_special_characters(self):
        """Test export with special characters in content."""
        exporter = DocxExporter()
        content = """Special characters test:

& ampersand
< less than
> greater than
" double quote
' single quote
© copyright
® registered
™ trademark"""

        result = exporter.export(content, "Special Chars Test")
        assert isinstance(result, bytes)
        assert len(result) > 0

    def test_export_with_unicode(self):
        """Test export with unicode characters."""
        exporter = DocxExporter()
        content = """Unicode test:

Chinese: 你好世界
Japanese: こんにちは
Korean: 안녕하세요
Arabic: مرحبا
"""

        result = exporter.export(content, "Unicode Test")
        assert isinstance(result, bytes)
        assert len(result) > 0

    def test_export_long_content(self):
        """Test export with very long content."""
        exporter = DocxExporter()
        # Generate ~100KB of content
        content = "Lorem ipsum dolor sit amet. " * 5000

        result = exporter.export(content, "Long Content Test")
        assert isinstance(result, bytes)
        assert len(result) > 0

    def test_export_empty_content(self):
        """Test export with empty content."""
        exporter = DocxExporter()
        result = exporter.export("", "Empty Test")

        assert isinstance(result, bytes)
        assert len(result) > 0

    def test_export_default_options(self):
        """Test that default options work correctly."""
        options = ExportOptions()

        assert options.font_name == "Times New Roman"
        assert options.font_size == 12
        assert options.margin_top == 1.0
        assert options.margin_bottom == 1.0
        assert options.margin_left == 1.0
        assert options.margin_right == 1.0
        assert options.line_spacing == 1.0
        assert options.include_letterhead is False
        assert options.include_page_numbers is True
        assert options.include_date is True

    def test_export_double_spacing(self):
        """Test export with double line spacing."""
        exporter = DocxExporter()
        options = ExportOptions(line_spacing=2.0)

        result = exporter.export("Test paragraph.\n\nAnother paragraph.", "Double Spacing Test", options)

        assert isinstance(result, bytes)
        assert len(result) > 0

    def test_mixed_formatting(self):
        """Test export with mixed formatting elements."""
        exporter = DocxExporter()
        content = """# Legal Demand Letter

**RE: Case #12345 - Smith v. Johnson**

## Background

This letter serves as a formal demand regarding the following matters:

1. Property damage assessment: **$50,000**
2. Medical expenses: **$25,000**
3. Lost wages: **$15,000**

### Summary of Damages

- Vehicle repair costs
- Hospital bills
- Physical therapy expenses

*Please respond within 30 days.*"""

        result = exporter.export(content, "Mixed Formatting Test")

        doc = Document(io.BytesIO(result))
        assert len(doc.paragraphs) > 10


if __name__ == "__main__":
    pytest.main([__file__, "-v"])
